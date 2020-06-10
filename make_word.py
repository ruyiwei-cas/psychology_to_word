from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.oxml.ns import qn
import pandas as pd


document = Document('./tree.docx')
document_save = Document()

df = pd.read_excel('./name_label.xlsx')
excel_vec = []
for index, row in df.iterrows():
    print(row)
    excel_vec.append(row)

string = '张三,01、1,02、1,03、1,0z4、1,05、1,06、1,07、1,08、1,09、1'
string_vec = string.split(',')

title = string_vec[0] + "心理测试结果"
document_save.add_heading(title, 0)

paragraphs = list(document.paragraphs)


string_index = 1
for index in range(len(paragraphs)):
    if paragraphs[index].text.find(string_vec[string_index]) != -1 or string_index < len(string_vec):
        string_index = string_index + 1
        if string_index > len(string_vec):
            break
        document_save.add_heading(paragraphs[index].text, level=2)
        # run = document_save.add_paragraph().add_run()
        # style = document.styles['Normal']
        # font = style.font
        # font.name = 'SimSum-ExtB'
        # font.size = Pt(15)
        # paragraph = document_save.add_paragraph(paragraphs[index + 1].text)
        run = document_save.add_paragraph().add_run(paragraphs[index + 1].text)
        run.font.name = u'宋体'
        r = run.element
        r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

save_name = './result/' + string_vec[0] + "心理测试结果.docx"
document_save.save(save_name)
